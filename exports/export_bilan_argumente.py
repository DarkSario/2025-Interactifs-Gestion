"""
Specialized export functions for "bilans argumentés" (argued/detailed financial reports).

This module provides functions to export detailed financial reports with arguments
and justifications in both PDF and Word formats. These reports are typically used
for end-of-fiscal-year reporting and include detailed breakdowns with commentary.
"""

import pandas as pd
from tkinter import filedialog, messagebox
from typing import Optional
from datetime import datetime


def export_bilan_argumente_pdf(file_path: Optional[str] = None):
    """
    Export a detailed "bilan argumenté" (argued balance sheet) to PDF format.
    
    This function generates a comprehensive financial report with detailed breakdowns
    by category, commentary, and analysis. It includes:
    - Revenue breakdown by source
    - Expense breakdown by category
    - Event-specific financial details
    - Commentary and justifications
    - Year-over-year comparisons (if applicable)
    
    Args:
        file_path: Optional output file path. If None, prompts user with file dialog.
        
    Example:
        >>> export_bilan_argumente_pdf('bilan_argumente_2023.pdf')
    """
    try:
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, 
            PageBreak, KeepTogether
        )
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
    except ImportError as e:
        error_msg = "reportlab is required for PDF export. Install it with: pip install reportlab"
        if messagebox:
            messagebox.showerror("Missing Dependency", error_msg)
        raise ImportError(error_msg) from e
    
    if file_path is None:
        file_path = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")],
            title="Export Bilan Argumenté (PDF)",
            initialfile=f"bilan_argumente_{datetime.now().year}.pdf"
        )
    
    if not file_path:
        return  # User cancelled
    
    try:
        from db.db import get_connection
        
        conn = get_connection()
        
        # Fetch detailed financial data
        # Events data
        events_df = pd.read_sql_query(
            "SELECT name, date, lieu FROM events ORDER BY date DESC", 
            conn
        )
        
        # Revenue by source
        recettes_events = pd.read_sql_query("""
            SELECT 
                source, 
                SUM(montant) as total,
                COUNT(*) as count
            FROM event_recettes 
            GROUP BY source
        """, conn)
        
        recettes_subventions = pd.read_sql_query("""
            SELECT 
                type_entite as source,
                SUM(montant) as total,
                COUNT(*) as count
            FROM dons_subventions
            GROUP BY type_entite
        """, conn)
        
        # Expenses by category
        depenses_events = pd.read_sql_query("""
            SELECT 
                categorie,
                SUM(montant) as total,
                COUNT(*) as count
            FROM event_depenses
            GROUP BY categorie
        """, conn)
        
        depenses_regulieres = pd.read_sql_query("""
            SELECT 
                categorie,
                SUM(montant) as total,
                COUNT(*) as count
            FROM depenses_regulieres
            GROUP BY categorie
        """, conn)
        
        depenses_diverses = pd.read_sql_query("""
            SELECT 
                categorie,
                SUM(montant) as total,
                COUNT(*) as count
            FROM depenses_diverses
            GROUP BY categorie
        """, conn)
        
        # Calculate totals
        total_recettes_events = recettes_events['total'].sum() if not recettes_events.empty else 0.0
        total_recettes_subventions = recettes_subventions['total'].sum() if not recettes_subventions.empty else 0.0
        total_recettes = total_recettes_events + total_recettes_subventions
        
        total_depenses_events = depenses_events['total'].sum() if not depenses_events.empty else 0.0
        total_depenses_regulieres = depenses_regulieres['total'].sum() if not depenses_regulieres.empty else 0.0
        total_depenses_diverses = depenses_diverses['total'].sum() if not depenses_diverses.empty else 0.0
        total_depenses = total_depenses_events + total_depenses_regulieres + total_depenses_diverses
        
        solde = total_recettes - total_depenses
        
        conn.close()
        
        # Create PDF document
        doc = SimpleDocTemplate(
            file_path, 
            pagesize=A4,
            rightMargin=30, 
            leftMargin=30, 
            topMargin=30, 
            bottomMargin=30
        )
        
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name='Justify',
            alignment=TA_JUSTIFY,
            fontSize=10,
            leading=14
        ))
        styles.add(ParagraphStyle(
            name='Center',
            alignment=TA_CENTER,
            fontSize=11
        ))
        
        elements = []
        
        # === TITLE PAGE ===
        elements.append(Spacer(1, 2*inch))
        elements.append(Paragraph(
            "<b>BILAN ARGUMENTÉ DE FIN D'EXERCICE</b>", 
            styles["Title"]
        ))
        elements.append(Spacer(1, 0.3*inch))
        elements.append(Paragraph(
            f"<i>Exercice clos le {datetime.now().strftime('%d/%m/%Y')}</i>",
            styles["Center"]
        ))
        elements.append(PageBreak())
        
        # === EXECUTIVE SUMMARY ===
        elements.append(Paragraph("<b>1. SYNTHÈSE GÉNÉRALE</b>", styles["Heading1"]))
        elements.append(Spacer(1, 12))
        
        summary_text = f"""
        Le présent bilan argumenté présente une analyse détaillée de l'exercice écoulé. 
        L'association a organisé {len(events_df)} événement(s) au cours de cette période. 
        Les résultats financiers font apparaître un solde {('positif' if solde >= 0 else 'négatif')} 
        de {abs(solde):.2f} €.
        """
        elements.append(Paragraph(summary_text, styles["Justify"]))
        elements.append(Spacer(1, 16))
        
        # Summary table
        summary_data = [
            ["<b>Poste</b>", "<b>Montant (€)</b>", "<b>%</b>"],
            ["Total Recettes", f"{total_recettes:.2f}", "100.00"],
            ["  - Recettes événements", f"{total_recettes_events:.2f}", 
             f"{(total_recettes_events/total_recettes*100 if total_recettes > 0 else 0):.2f}"],
            ["  - Subventions/Dons", f"{total_recettes_subventions:.2f}",
             f"{(total_recettes_subventions/total_recettes*100 if total_recettes > 0 else 0):.2f}"],
            ["", "", ""],
            ["Total Dépenses", f"{total_depenses:.2f}", 
             f"{(total_depenses/total_recettes*100 if total_recettes > 0 else 0):.2f}"],
            ["  - Dépenses événements", f"{total_depenses_events:.2f}",
             f"{(total_depenses_events/total_recettes*100 if total_recettes > 0 else 0):.2f}"],
            ["  - Dépenses régulières", f"{total_depenses_regulieres:.2f}",
             f"{(total_depenses_regulieres/total_recettes*100 if total_recettes > 0 else 0):.2f}"],
            ["  - Dépenses diverses", f"{total_depenses_diverses:.2f}",
             f"{(total_depenses_diverses/total_recettes*100 if total_recettes > 0 else 0):.2f}"],
            ["", "", ""],
            ["<b>SOLDE</b>", f"<b>{solde:.2f}</b>", 
             f"<b>{(solde/total_recettes*100 if total_recettes > 0 else 0):.2f}</b>"],
        ]
        
        summary_table = Table(summary_data, colWidths=[3*inch, 1.5*inch, 0.8*inch])
        summary_table.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.7, colors.grey),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BACKGROUND", (0, 10), (-1, 10), colors.HexColor("#ffffcc")),
            ("FONTNAME", (0, 10), (-1, 10), "Helvetica-Bold"),
            ("ROWBACKGROUNDS", (0, 1), (-1, 9), [colors.white, colors.HexColor("#f0f0f0")]),
        ]))
        
        elements.append(summary_table)
        elements.append(Spacer(1, 20))
        
        # === REVENUE ANALYSIS ===
        elements.append(Paragraph("<b>2. ANALYSE DES RECETTES</b>", styles["Heading1"]))
        elements.append(Spacer(1, 12))
        
        recettes_text = f"""
        Les recettes de l'exercice s'élèvent à {total_recettes:.2f} €, réparties entre 
        les revenus générés par les événements ({total_recettes_events:.2f} €) et les 
        subventions/dons reçus ({total_recettes_subventions:.2f} €).
        """
        elements.append(Paragraph(recettes_text, styles["Justify"]))
        elements.append(Spacer(1, 12))
        
        # Revenue breakdown table
        if not recettes_events.empty:
            elements.append(Paragraph("<b>2.1. Recettes par événements (par source)</b>", styles["Heading2"]))
            elements.append(Spacer(1, 8))
            
            recettes_table_data = [["<b>Source</b>", "<b>Nombre</b>", "<b>Total (€)</b>"]]
            for _, row in recettes_events.iterrows():
                recettes_table_data.append([
                    str(row['source']),
                    str(int(row['count'])),
                    f"{row['total']:.2f}"
                ])
            
            recettes_table = Table(recettes_table_data, colWidths=[3*inch, 1*inch, 1.5*inch])
            recettes_table.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e0e0e0")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("ALIGN", (1, 1), (-1, -1), "RIGHT"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.lightgrey]),
            ]))
            elements.append(recettes_table)
            elements.append(Spacer(1, 16))
        
        # === EXPENSE ANALYSIS ===
        elements.append(Paragraph("<b>3. ANALYSE DES DÉPENSES</b>", styles["Heading1"]))
        elements.append(Spacer(1, 12))
        
        depenses_text = f"""
        Les dépenses totales de l'exercice s'élèvent à {total_depenses:.2f} €, 
        comprenant les dépenses liées aux événements, les dépenses régulières de 
        fonctionnement, et les dépenses diverses.
        """
        elements.append(Paragraph(depenses_text, styles["Justify"]))
        elements.append(Spacer(1, 12))
        
        # Expense breakdown tables
        if not depenses_events.empty:
            elements.append(Paragraph("<b>3.1. Dépenses événements (par catégorie)</b>", styles["Heading2"]))
            elements.append(Spacer(1, 8))
            
            depenses_table_data = [["<b>Catégorie</b>", "<b>Nombre</b>", "<b>Total (€)</b>"]]
            for _, row in depenses_events.iterrows():
                depenses_table_data.append([
                    str(row['categorie']),
                    str(int(row['count'])),
                    f"{row['total']:.2f}"
                ])
            
            depenses_table = Table(depenses_table_data, colWidths=[3*inch, 1*inch, 1.5*inch])
            depenses_table.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e0e0e0")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("ALIGN", (1, 1), (-1, -1), "RIGHT"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.lightgrey]),
            ]))
            elements.append(depenses_table)
            elements.append(Spacer(1, 16))
        
        # === CONCLUSION ===
        elements.append(Paragraph("<b>4. CONCLUSION</b>", styles["Heading1"]))
        elements.append(Spacer(1, 12))
        
        conclusion_text = f"""
        L'exercice se clôture sur un résultat {('excédentaire' if solde >= 0 else 'déficitaire')} 
        de {abs(solde):.2f} €. Cette situation témoigne de {'la bonne santé financière de l\'association' if solde >= 0 else 'difficultés qu\'il conviendra d\'analyser en détail'}. 
        Les perspectives pour l'exercice à venir devront tenir compte de ces résultats pour 
        ajuster la stratégie financière de l'association.
        """
        elements.append(Paragraph(conclusion_text, styles["Justify"]))
        elements.append(Spacer(1, 20))
        
        # Footer
        elements.append(Spacer(1, 0.5*inch))
        elements.append(Paragraph(
            f"<i>Document généré automatiquement le {datetime.now().strftime('%d/%m/%Y à %H:%M')}</i>",
            styles["Center"]
        ))
        
        # Build PDF
        doc.build(elements)
        
        if messagebox:
            messagebox.showinfo(
                "Export Success", 
                f"Bilan argumenté (PDF) exported successfully to:\n{file_path}"
            )
            
    except Exception as e:
        if messagebox:
            messagebox.showerror("Export Error", f"Failed to generate bilan argumenté:\n{str(e)}")
        raise


def export_bilan_argumente_word(file_path: Optional[str] = None):
    """
    Export a detailed "bilan argumenté" (argued balance sheet) to Word format.
    
    This function generates a comprehensive financial report in Microsoft Word format
    (.docx) with the same detailed breakdowns as the PDF version but in an editable format.
    
    Args:
        file_path: Optional output file path. If None, prompts user with file dialog.
        
    Raises:
        ImportError: If python-docx is not installed
        
    Example:
        >>> export_bilan_argumente_word('bilan_argumente_2023.docx')
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as e:
        error_msg = "python-docx is required for Word export. Install it with: pip install python-docx"
        if messagebox:
            messagebox.showerror("Missing Dependency", error_msg)
        raise ImportError(error_msg) from e
    
    if file_path is None:
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")],
            title="Export Bilan Argumenté (Word)",
            initialfile=f"bilan_argumente_{datetime.now().year}.docx"
        )
    
    if not file_path:
        return  # User cancelled
    
    try:
        from db.db import get_connection
        
        conn = get_connection()
        
        # Fetch detailed financial data (same as PDF version)
        events_df = pd.read_sql_query(
            "SELECT name, date, lieu FROM events ORDER BY date DESC", 
            conn
        )
        
        recettes_events = pd.read_sql_query("""
            SELECT source, SUM(montant) as total, COUNT(*) as count
            FROM event_recettes GROUP BY source
        """, conn)
        
        recettes_subventions = pd.read_sql_query("""
            SELECT type_entite as source, SUM(montant) as total, COUNT(*) as count
            FROM dons_subventions GROUP BY type_entite
        """, conn)
        
        depenses_events = pd.read_sql_query("""
            SELECT categorie, SUM(montant) as total, COUNT(*) as count
            FROM event_depenses GROUP BY categorie
        """, conn)
        
        depenses_regulieres = pd.read_sql_query("""
            SELECT categorie, SUM(montant) as total, COUNT(*) as count
            FROM depenses_regulieres GROUP BY categorie
        """, conn)
        
        depenses_diverses = pd.read_sql_query("""
            SELECT categorie, SUM(montant) as total, COUNT(*) as count
            FROM depenses_diverses GROUP BY categorie
        """, conn)
        
        # Calculate totals
        total_recettes_events = recettes_events['total'].sum() if not recettes_events.empty else 0.0
        total_recettes_subventions = recettes_subventions['total'].sum() if not recettes_subventions.empty else 0.0
        total_recettes = total_recettes_events + total_recettes_subventions
        
        total_depenses_events = depenses_events['total'].sum() if not depenses_events.empty else 0.0
        total_depenses_regulieres = depenses_regulieres['total'].sum() if not depenses_regulieres.empty else 0.0
        total_depenses_diverses = depenses_diverses['total'].sum() if not depenses_diverses.empty else 0.0
        total_depenses = total_depenses_events + total_depenses_regulieres + total_depenses_diverses
        
        solde = total_recettes - total_depenses
        
        conn.close()
        
        # Create Word document
        doc = Document()
        
        # === TITLE PAGE ===
        title = doc.add_heading("BILAN ARGUMENTÉ DE FIN D'EXERCICE", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph(f"Exercice clos le {datetime.now().strftime('%d/%m/%Y')}")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_format = subtitle.runs[0].font
        subtitle_format.italic = True
        subtitle_format.size = Pt(12)
        
        doc.add_page_break()
        
        # === EXECUTIVE SUMMARY ===
        doc.add_heading("1. SYNTHÈSE GÉNÉRALE", 1)
        
        summary_text = (
            f"Le présent bilan argumenté présente une analyse détaillée de l'exercice écoulé. "
            f"L'association a organisé {len(events_df)} événement(s) au cours de cette période. "
            f"Les résultats financiers font apparaître un solde "
            f"{'positif' if solde >= 0 else 'négatif'} de {abs(solde):.2f} €."
        )
        doc.add_paragraph(summary_text)
        
        # Summary table
        table = doc.add_table(rows=11, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Poste"
        header_cells[1].text = "Montant (€)"
        header_cells[2].text = "%"
        
        # Data rows
        table.rows[1].cells[0].text = "Total Recettes"
        table.rows[1].cells[1].text = f"{total_recettes:.2f}"
        table.rows[1].cells[2].text = "100.00"
        
        table.rows[2].cells[0].text = "  - Recettes événements"
        table.rows[2].cells[1].text = f"{total_recettes_events:.2f}"
        table.rows[2].cells[2].text = f"{(total_recettes_events/total_recettes*100 if total_recettes > 0 else 0):.2f}"
        
        table.rows[3].cells[0].text = "  - Subventions/Dons"
        table.rows[3].cells[1].text = f"{total_recettes_subventions:.2f}"
        table.rows[3].cells[2].text = f"{(total_recettes_subventions/total_recettes*100 if total_recettes > 0 else 0):.2f}"
        
        table.rows[5].cells[0].text = "Total Dépenses"
        table.rows[5].cells[1].text = f"{total_depenses:.2f}"
        table.rows[5].cells[2].text = f"{(total_depenses/total_recettes*100 if total_recettes > 0 else 0):.2f}"
        
        table.rows[6].cells[0].text = "  - Dépenses événements"
        table.rows[6].cells[1].text = f"{total_depenses_events:.2f}"
        table.rows[6].cells[2].text = f"{(total_depenses_events/total_recettes*100 if total_recettes > 0 else 0):.2f}"
        
        table.rows[7].cells[0].text = "  - Dépenses régulières"
        table.rows[7].cells[1].text = f"{total_depenses_regulieres:.2f}"
        table.rows[7].cells[2].text = f"{(total_depenses_regulieres/total_recettes*100 if total_recettes > 0 else 0):.2f}"
        
        table.rows[8].cells[0].text = "  - Dépenses diverses"
        table.rows[8].cells[1].text = f"{total_depenses_diverses:.2f}"
        table.rows[8].cells[2].text = f"{(total_depenses_diverses/total_recettes*100 if total_recettes > 0 else 0):.2f}"
        
        table.rows[10].cells[0].text = "SOLDE"
        table.rows[10].cells[1].text = f"{solde:.2f}"
        table.rows[10].cells[2].text = f"{(solde/total_recettes*100 if total_recettes > 0 else 0):.2f}"
        
        # Make last row bold
        for cell in table.rows[10].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        doc.add_paragraph()  # Spacing
        
        # === REVENUE ANALYSIS ===
        doc.add_heading("2. ANALYSE DES RECETTES", 1)
        
        recettes_text = (
            f"Les recettes de l'exercice s'élèvent à {total_recettes:.2f} €, réparties entre "
            f"les revenus générés par les événements ({total_recettes_events:.2f} €) et les "
            f"subventions/dons reçus ({total_recettes_subventions:.2f} €)."
        )
        doc.add_paragraph(recettes_text)
        
        if not recettes_events.empty:
            doc.add_heading("2.1. Recettes par événements (par source)", 2)
            
            recettes_table = doc.add_table(rows=len(recettes_events) + 1, cols=3)
            recettes_table.style = 'Light Grid Accent 1'
            
            # Header
            hdr_cells = recettes_table.rows[0].cells
            hdr_cells[0].text = "Source"
            hdr_cells[1].text = "Nombre"
            hdr_cells[2].text = "Total (€)"
            
            # Data
            for i, (_, row) in enumerate(recettes_events.iterrows(), 1):
                row_cells = recettes_table.rows[i].cells
                row_cells[0].text = str(row['source'])
                row_cells[1].text = str(int(row['count']))
                row_cells[2].text = f"{row['total']:.2f}"
        
        # === EXPENSE ANALYSIS ===
        doc.add_heading("3. ANALYSE DES DÉPENSES", 1)
        
        depenses_text = (
            f"Les dépenses totales de l'exercice s'élèvent à {total_depenses:.2f} €, "
            f"comprenant les dépenses liées aux événements, les dépenses régulières de "
            f"fonctionnement, et les dépenses diverses."
        )
        doc.add_paragraph(depenses_text)
        
        if not depenses_events.empty:
            doc.add_heading("3.1. Dépenses événements (par catégorie)", 2)
            
            depenses_table = doc.add_table(rows=len(depenses_events) + 1, cols=3)
            depenses_table.style = 'Light Grid Accent 1'
            
            # Header
            hdr_cells = depenses_table.rows[0].cells
            hdr_cells[0].text = "Catégorie"
            hdr_cells[1].text = "Nombre"
            hdr_cells[2].text = "Total (€)"
            
            # Data
            for i, (_, row) in enumerate(depenses_events.iterrows(), 1):
                row_cells = depenses_table.rows[i].cells
                row_cells[0].text = str(row['categorie'])
                row_cells[1].text = str(int(row['count']))
                row_cells[2].text = f"{row['total']:.2f}"
        
        # === CONCLUSION ===
        doc.add_heading("4. CONCLUSION", 1)
        
        conclusion_text = (
            f"L'exercice se clôture sur un résultat "
            f"{'excédentaire' if solde >= 0 else 'déficitaire'} de {abs(solde):.2f} €. "
            f"Cette situation témoigne de "
            f"{'la bonne santé financière de l\'association' if solde >= 0 else 'difficultés qu\'il conviendra d\'analyser en détail'}. "
            f"Les perspectives pour l'exercice à venir devront tenir compte de ces résultats pour "
            f"ajuster la stratégie financière de l'association."
        )
        doc.add_paragraph(conclusion_text)
        
        # Footer
        footer_p = doc.add_paragraph()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_p.add_run(
            f"Document généré automatiquement le {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
        )
        footer_run.font.italic = True
        footer_run.font.size = Pt(9)
        
        # Save document
        doc.save(file_path)
        
        if messagebox:
            messagebox.showinfo(
                "Export Success", 
                f"Bilan argumenté (Word) exported successfully to:\n{file_path}"
            )
            
    except Exception as e:
        if messagebox:
            messagebox.showerror("Export Error", f"Failed to generate bilan argumenté:\n{str(e)}")
        raise
